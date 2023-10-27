import os  # Import the os module here
from django.conf import settings
from django.shortcuts import render
from django.templatetags.static import static
from django.contrib.staticfiles.finders import find
from masque_EME.forms import pre40
from pdfminer.high_level import extract_text
import re
from docx import Document
from django.conf import settings
from django.contrib.auth.decorators import login_required
import zipfile
import shutil
from datetime import datetime


def replace_tag(paragraphs, tag, new_text):
    for paragraph in paragraphs:
        if tag in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(tag, new_text)


def replace_tag_in_docx(file_path, tag_dict,doc_name):
    doc = Document(file_path)
    
    # Replace in paragraphs
    for tag in tag_dict:
        replace_tag(doc.paragraphs, tag, tag_dict[tag])
    
    # Replace in tables
    for tag in tag_dict:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_tag(cell.paragraphs, tag, tag_dict[tag])
                
    # Replace in headers
    for tag in tag_dict:
        for section in doc.sections:
            for header in section.header.paragraphs:
                replace_tag(section.header.paragraphs, tag, tag_dict[tag])
                
            # Replace in footers
            for footer in section.footer.paragraphs:
                replace_tag(section.footer.paragraphs, tag, tag_dict[tag])
            
    # Define the path to save the modified document
    new_file_name = doc_name + '_' + tag_dict['N_MARCHE'] + '_'+tag_dict['N_COMMANDE']+'_'+tag_dict['BENEFICIARY_NOM'].replace(" ","_")+'_'+tag_dict['DATE_START'].replace("/","_")+'.docx'
    media_root = settings.MEDIA_ROOT
    save_path = os.path.join(media_root, new_file_name)
    
    # Save the modified document
    doc.save(save_path)

    # Return the URL
    url = f'{settings.MEDIA_URL}{new_file_name}'
    url_from_app = f'{new_file_name}'
    dict = {'url':url,'url_from_app':url_from_app}

    return dict



def extract_information(text):
    info_dict = {}
    emails = re.findall(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+(?:\.[a-zA-Z0-9-.]+)?", text)
    phones = re.findall(r"\b\d{8,15}\b", text)

    # Regular expressions for each information to be captured
    info_dict["N° de Marché"] = re.search(r"N° marché : (\d+)", text).group(1) if re.search(r"N° marché : (\d+)", text) else "Not found"
    info_dict["N° du lot"] = re.search(r"N° du lot : (\d+)", text).group(1) if re.search(r"N° du lot : (\d+)", text) else "Not found"
    info_dict["N° de commande"] = re.search(r"N° commande :\n\n(\w+)", text).group(1) if re.search(r"N° commande :\n\n(\w+)", text) else "Not found"
    
    # Capturing start date and reformatting
    start_date_match = re.search(r"Prestation du (\d{2})/(\d{2})/(\d{4})", text)
    if start_date_match:
        info_dict["Prestation réalisée du"] = f"{start_date_match.group(1)}/{start_date_match.group(2)}/{start_date_match.group(3)}"
        info_dict["date_start_format"] = f"|{start_date_match.group(1)[0]}|{start_date_match.group(1)[1]}| / |{start_date_match.group(2)[0]}|{start_date_match.group(2)[1]}| / |{start_date_match.group(3)[2]}|{start_date_match.group(3)[3]}|"
    else:
        info_dict["Prestation réalisée du"] = "Not found"
        info_dict["date_start_format"] = "Not found"

    # Capturing end date and reformatting
    end_date_match = re.search(r"au (\d{2})/(\d{2})/(\d{4})", text)
    if end_date_match:
        info_dict["Prestation réalisée au"] = f"{end_date_match.group(1)}/{end_date_match.group(2)}/{end_date_match.group(3)}"
        info_dict["date_end_format"] = f"|{end_date_match.group(1)[0]}|{end_date_match.group(1)[1]}| / |{end_date_match.group(2)[0]}|{end_date_match.group(2)[1]}| / |{end_date_match.group(3)[2]}|{end_date_match.group(3)[3]}|"
    else:
        info_dict["Prestation réalisée au"] = "Not found"
        info_dict["date_end_format"] = "Not found"

    info_dict["Bénéficiaire > Identifiant"] = re.search(r"Identifiant N° : (\w+)", text).group(1) if re.search(r"Identifiant N° : (\w+)", text) else "Not found"

    beneficiary_name_match = re.search(r"Nom, prénom : (.*?)Nom :", text, re.DOTALL)
    if beneficiary_name_match:
        beneficiary_name = beneficiary_name_match.group(1).strip()  # Removing spaces before and after the string
        info_dict["Bénéficiaire > Nom, Prénom"] = beneficiary_name
    else:
        info_dict["Bénéficiaire > Nom, Prénom"] = "Not found"

    info_dict["Bénéficiaire > tel"] = phones[0] if phones else "Not found"

    # Formatting and assigning the first phone number found
    if phones:
        formatted_phone = '|' + '|'.join(phones[0]) + '|'
        info_dict["b_tel_format"] = formatted_phone
    else:
        info_dict["b_tel_format"] = "Not found"

    info_dict["Bénéficiaire > mail"] = emails[0] if emails else "Not found"
    
    organisme_nom_match = re.search(r"Nom : (.+?)\n\nIdentifiant N°", text, re.DOTALL)
    if organisme_nom_match:
        info_dict["Organisme > Nom"] = organisme_nom_match.group(1).strip()
    else:
        info_dict["Organisme > Nom"] = "Not found"

    lieu_realisation_match = re.search(r"100% à distance", text)

    if lieu_realisation_match:  # If "100% à distance" is found
        info_dict["Organisme > lieu de réalisation"] = lieu_realisation_match.group(0)
    else:  # If "100% à distance" is not found
        # Trying to find text starting with "BGE ADIL -" up until "Tél"
        lieu_realisation_match = re.search(r"BGE ADIL -.*?(?=Tél)", text, re.DOTALL)
        if lieu_realisation_match:
            lieu_realisation = lieu_realisation_match.group(0).strip()
            info_dict["Organisme > lieu de réalisation"] = lieu_realisation
        else:
            info_dict["Organisme > lieu de réalisation"] = "Not found"

    info_dict["Organisme > Tel"] = "0145805155" # If there is a specific pattern to capture, add it here.
    info_dict["Organisme > Mail"] = emails[1] if len(emails) > 1 else "Not found"
    
    # Pole Emploi
    pole_emploi_match = re.search(r"Pôle emploi de\s*:([^\n]*)", text)
    if pole_emploi_match:
        pole_emploi = pole_emploi_match.group(1).strip()  # Stripping spaces around the string
        info_dict["Correspondant > Pole emploi de"] = pole_emploi
    else:
        info_dict["Correspondant > Pole emploi de"] = "Not found"



    # Conseiller PE
    correspondent_name_match = re.search(r"Correspondant-e Pôle emploi\n\nNom, prénom :.*?(\w+[\s\n]*\w+).*?Pôle emploi de :", text, re.DOTALL)
    if correspondent_name_match:
        correspondent_name = correspondent_name_match.group(1).replace('\n', ' ').strip()  # Removing spaces and newline characters around and within the string
        info_dict["Correspondant > Nom, prénom"] = correspondent_name
    else:
        info_dict["Correspondant > Nom, prénom"] = "Not found"
           
    return info_dict

@login_required
def index(request):
    file_url = None
    pre40_url = None
    pre20_url = None    
    pre40f = None
    pre20f = None
    extracted_info = None
    extracted_text = None  
    tagged_info = None
    zip_file_url = None
    saved_data = request.user.action_data or {}
    message = None

    initial_data = {
        'mail_conseiller': request.user.email,
        'conseiller': request.user.last_name + " " + request.user.first_name,
        'rdv_diagnostic': saved_data.get('rdv_diagnostic', ''),
        'atelier_1': saved_data.get('atelier_1', ''),
        'point_etape_1': saved_data.get('point_etape_1', ''),
        'atelier_2': saved_data.get('atelier_2', ''),
        'webinaire': saved_data.get('webinaire', ''),
        'point_etape_2': saved_data.get('point_etape_2', ''),
        'atelier_3': saved_data.get('atelier_3', ''),
        'point_etape_3': saved_data.get('point_etape_3', ''),
        'rdv_intermediaire': saved_data.get('rdv_intermediaire', ''),
        'atelier_4': saved_data.get('atelier_4', ''),
        'rdv_bilan': saved_data.get('rdv_bilan', ''),
    }

    # Mapping of info_dict keys to tags
    tag_mapping = {
        "N° de Marché": "N_MARCHE",
        "N° du lot": "N_LOT",
        "N° de commande": "N_COMMANDE",
        "date_start_format":"DSFORMAT",
        "date_end_format":"DEFORMAT",
        "Prestation réalisée du": "DATE_START",
        "Prestation réalisée au": "DATE_END",
        "Bénéficiaire > Identifiant": "BENEFICIARY_ID",
        "Bénéficiaire > Nom, Prénom":"BENEFICIARY_NOM",
        "Bénéficiaire > tel": "BENEFICIARY_TEL",
        "b_tel_format":"BTELF",
        "Bénéficiaire > mail": "BENEFICIARY_MAIL",
        "Organisme > Nom": "ORGANISM_NOM",
        "Organisme > lieu de réalisation": "ORGANISM_LIEU",
        "Organisme > Tel": "ORGANISM_TEL",
        "Organisme > Mail": "ORGANISM_MAIL",
        "Correspondant > Pole emploi de": "CORRESPONDANT_POLE_EMPLOI",
        "Correspondant > Nom, prénom": "CORRESPONDANT_NOM",
    }

    if request.method == "POST":
        form = pre40(request.POST, request.FILES)
        
        if form.is_valid():
            conseiller= form.cleaned_data.get('conseiller')  # Getting the content of 'conseiller' field
            mail_conseiller= form.cleaned_data.get('mail_conseiller')  # Getting the content of 'conseiller' field
            # Capturing 'rdv_bilan' date
            date_str = form.cleaned_data['rdv_bilan']
            
            try:
                rdv_bilan = datetime.strptime(date_str, '%d/%m/%Y')
            except (ValueError, TypeError):
                rdv_bilan = None


            # Formatting 'rdv_bilan' date
            if rdv_bilan:  # Check if 'rdv_bilan' is not None or empty
                rdv_bilan_format = rdv_bilan.strftime("|%d| / |%m| / |%y|")
            else:
                rdv_bilan_format = ""



            planification_data = {
                'conseiller': form.cleaned_data['conseiller'],
                'mail_conseiller': form.cleaned_data['mail_conseiller'],
                'rdv_diagnostic': form.cleaned_data['rdv_diagnostic'],
                'atelier_1': form.cleaned_data['atelier_1'],
                'point_etape_1': form.cleaned_data['point_etape_1'],
                'atelier_2': form.cleaned_data['atelier_2'],
                'point_etape_2': form.cleaned_data['point_etape_2'],
                'webinaire': form.cleaned_data['webinaire'],
                'atelier_3': form.cleaned_data['atelier_3'],
                'point_etape_3': form.cleaned_data['point_etape_3'],
                'rdv_intermediaire': form.cleaned_data['rdv_intermediaire'],
                'atelier_4': form.cleaned_data['atelier_4'],
                'rdv_bilan': form.cleaned_data['rdv_bilan'],
            }

            request.user.action_data = planification_data
            request.user.save()

            

            instance = form.save()
            file_url = instance.pdf_file.url
            file_path = instance.pdf_file.path

            # Extract text from the uploaded PDF
            extracted_text = extract_text(file_path)
            
            # Extract specific information from the text
            extracted_info = extract_information(extracted_text)
            
            # Creating a new dictionary with tags as keys
            tagged_info = {tag_mapping[key]: value for key, value in extracted_info.items() if key in tag_mapping}

            # Adding 'conseiller' and 'mail_conseiller' values to tagged_info
            tagged_info['MAIL_CONSEILLER'] = mail_conseiller
            tagged_info['CONSEILLER'] = conseiller

            tagged_info['RDV_BILAN_FORMAT'] = rdv_bilan_format


             # Adding 'planification_data' values to tagged_info with prefix PLANIF_
            for key, value in planification_data.items():
                tagged_info[f'PLANIF_{key.upper()}'] = value

            if "distance" in tagged_info["ORGANISM_LIEU"]:
                pre40f = replace_tag_in_docx('static/Skoll/docx/PRE40_dist.docx',tagged_info,'PRE40')
                message = "Modalité à distance"
            else:
                pre40f = replace_tag_in_docx('static/Skoll/docx/PRE40_pres.docx',tagged_info,'PRE40')
                message = "Modalité en présence"


            pre20f = replace_tag_in_docx('static/Skoll/docx/PRE20.docx',tagged_info,'PRE20')

            pre20_url = pre20f['url']
            pre40_url = pre40f['url']

            # Creating the zip file
            zip_file_name = tagged_info['N_COMMANDE'] + "_" + tagged_info['BENEFICIARY_NOM'] + '.zip'

            zip_file_path = os.path.join(settings.MEDIA_ROOT, zip_file_name)
            with zipfile.ZipFile(zip_file_path, 'w') as zf:
                zf.write(os.path.join(settings.MEDIA_ROOT, pre40f['url_from_app']), pre40f['url_from_app'])
                zf.write(os.path.join(settings.MEDIA_ROOT, pre20f['url_from_app']), pre20f['url_from_app'])
                zf.write(file_path, os.path.basename(file_path))
            
            zip_file_url = f'{settings.MEDIA_URL}{zip_file_name}'

            saved_data = request.user.action_data or {}
            initial_data = {
                'mail_conseiller': request.user.email,
                'conseiller': request.user.last_name + " " + request.user.first_name,
                'rdv_diagnostic': saved_data.get('rdv_diagnostic', ''),
                'atelier_1': saved_data.get('atelier_1', ''),
                'point_etape_1': saved_data.get('point_etape_1', ''),
                'atelier_2': saved_data.get('atelier_2', ''),
                'webinaire': saved_data.get('webinaire', ''),
                'point_etape_2': saved_data.get('point_etape_2', ''),
                'atelier_3': saved_data.get('atelier_3', ''),
                'point_etape_3': saved_data.get('point_etape_3', ''),
                'rdv_intermediaire': saved_data.get('rdv_intermediaire', ''),
                'atelier_4': saved_data.get('atelier_4', ''),
                'rdv_bilan': saved_data.get('rdv_bilan', ''),
            }

            form = pre40(initial=initial_data)  # Reset the form after saving
            

    else:
        form = pre40(initial=initial_data)

    context = {
        "form": form,
        "file_url": file_url,
        "extracted_info": tagged_info,
        "extracted_text": extracted_text,
        "PRE40_url": pre40_url,
        "PRE20_url": pre20_url,
        "zip_file": zip_file_url,
        "message":message,
    }

    return render(request, 'masque_EME/index.html', context)